from __future__ import annotations

import re
from collections import defaultdict
from pathlib import Path
from xml.sax.saxutils import escape
from zipfile import ZIP_DEFLATED, ZipFile

import pandas as pd
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

try:
    from openpyxl.cell.rich_text import CellRichText, TextBlock
    from openpyxl.cell.text import InlineFont

    _OPENPYXL_RICH_TEXT_AVAILABLE = True
except Exception:
    CellRichText = None
    TextBlock = None
    InlineFont = None
    _OPENPYXL_RICH_TEXT_AVAILABLE = False

from phonetic_toolbox.core.transcription.phonology_induction import (
    PhonologyInductionParser,
)
from phonetic_toolbox.models.phonology_models import (
    ParsedPhonologyRow,
    PhonologyAnalysisResult,
    PhonologyInputRow,
    PhonologyOutputResult,
)


class _FallbackFont:
    def __init__(self, run: "_FallbackRun"):
        self._run = run

    @property
    def subscript(self) -> bool:
        return self._run.subscript

    @subscript.setter
    def subscript(self, value: bool):
        self._run.subscript = bool(value)


class _FallbackRun:
    def __init__(self, text: str):
        self.text = text
        self.subscript = False
        self.bold = False
        self._font = _FallbackFont(self)

    @property
    def font(self) -> _FallbackFont:
        return self._font


class _FallbackParagraph:
    def __init__(self):
        self.runs: list[_FallbackRun] = []

    def add_run(self, text: str = "") -> _FallbackRun:
        run = _FallbackRun(text)
        self.runs.append(run)
        return run


class _FallbackDocument:
    def __init__(self):
        self.paragraphs: list[_FallbackParagraph] = []

    def add_heading(self, text: str, level: int = 1) -> _FallbackParagraph:
        paragraph = self.add_paragraph()
        run = paragraph.add_run(text)
        run.bold = True
        return paragraph

    def add_paragraph(self, text: str = "") -> _FallbackParagraph:
        paragraph = _FallbackParagraph()
        self.paragraphs.append(paragraph)
        if text:
            paragraph.add_run(text)
        return paragraph

    def save(self, output_path: Path):
        output_path = Path(output_path)
        document_xml = self._build_document_xml()
        content_types_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '<Override PartName="/word/document.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/>'
            "</Types>"
        )
        rels_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
            '<Relationship Id="rId1" '
            'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
            'Target="word/document.xml"/>'
            "</Relationships>"
        )
        document_rels_xml = (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"/>'
        )
        with ZipFile(output_path, "w", compression=ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", content_types_xml)
            zf.writestr("_rels/.rels", rels_xml)
            zf.writestr("word/document.xml", document_xml)
            zf.writestr("word/_rels/document.xml.rels", document_rels_xml)

    def _build_document_xml(self) -> str:
        paragraph_xml = "".join(self._render_paragraph(paragraph) for paragraph in self.paragraphs)
        return (
            '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
            '<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            f"<w:body>{paragraph_xml}<w:sectPr/></w:body>"
            "</w:document>"
        )

    def _render_paragraph(self, paragraph: _FallbackParagraph) -> str:
        if not paragraph.runs:
            return "<w:p/>"
        runs_xml = "".join(self._render_run(run) for run in paragraph.runs)
        return f"<w:p>{runs_xml}</w:p>"

    def _render_run(self, run: _FallbackRun) -> str:
        text = escape(run.text)
        if not run.subscript and not run.bold:
            return f'<w:r><w:t xml:space="preserve">{text}</w:t></w:r>'
        props = []
        if run.bold:
            props.append("<w:b/>")
        if run.subscript:
            props.append('<w:vertAlign w:val="subscript"/>')
        rpr = f"<w:rPr>{''.join(props)}</w:rPr>"
        return f'<w:r>{rpr}<w:t xml:space="preserve">{text}</w:t></w:r>'


class PhonologyInductionService:
    INITIAL_MANNER_ORDER = [
        "鼻音",
        "塞音",
        "塞擦音",
        "擦音",
        "近音",
        "边音",
        "闪音",
        "颤音",
        "其他",
    ]
    INITIAL_PLACE_ORDER = [
        "双唇",
        "唇齿",
        "舌尖前",
        "舌尖中",
        "舌尖后",
        "舌叶",
        "卷舌",
        "龈后",
        "龈腭",
        "硬腭",
        "软腭",
        "小舌",
        "咽",
        "声门",
        "其他",
    ]

    def __init__(self):
        self._parser = PhonologyInductionParser()

    def load_rows(self, file_path: str, skip_first_row: bool = False) -> list[PhonologyInputRow]:
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        suffix = path.suffix.lower()
        if suffix in {".xlsx", ".xls"}:
            return self._load_rows_from_excel(path, skip_first_row=skip_first_row)
        if suffix == ".csv":
            return self._load_rows_from_csv(path, skip_first_row=skip_first_row)
        return self._load_rows_from_text(path, skip_first_row=skip_first_row)

    def analyze(
        self,
        rows: list[PhonologyInputRow],
        consonant_only_as_zero_initial: bool = True,
    ) -> PhonologyAnalysisResult:
        parsed_rows: list[ParsedPhonologyRow] = []
        unique_ipa: set[str] = set()
        for row in rows:
            parsed = self._parser.parse(
                row.ipa,
                consonant_only_as_zero_initial=consonant_only_as_zero_initial,
            )
            parsed_rows.append(
                ParsedPhonologyRow(
                    character=row.character,
                    ipa=row.ipa,
                    note=row.note,
                    initial=parsed.initial,
                    final=parsed.final,
                    tone_value=parsed.tone,
                )
            )
            unique_ipa.add(row.ipa)
        unique_initials = self._sort_symbols({row.initial for row in parsed_rows})
        unique_initials = self._sort_initials(set(unique_initials))
        unique_finals = self._sort_symbols({row.final for row in parsed_rows})
        unique_tones = self._sort_tones({row.tone_value for row in parsed_rows})
        return PhonologyAnalysisResult(
            rows=parsed_rows,
            unique_initials=unique_initials,
            unique_finals=unique_finals,
            unique_tones=unique_tones,
            unique_ipa=sorted(unique_ipa),
        )

    def find_single_consonant_rows(
        self, rows: list[PhonologyInputRow]
    ) -> list[PhonologyInputRow]:
        return [row for row in rows if self._parser.is_single_consonant_syllable(row.ipa)]

    def apply_symbol_aliases(
        self,
        analysis: PhonologyAnalysisResult,
        initial_merge_map: dict[str, str],
        final_merge_map: dict[str, str],
    ) -> PhonologyAnalysisResult:
        def resolve(value: str, merge_map: dict[str, str]) -> str:
            current = value
            visited: set[str] = set()
            while current in merge_map and current not in visited:
                visited.add(current)
                current = merge_map[current]
            return current

        remapped_rows: list[ParsedPhonologyRow] = []
        for row in analysis.rows:
            remapped_rows.append(
                ParsedPhonologyRow(
                    character=row.character,
                    ipa=row.ipa,
                    note=row.note,
                    initial=resolve(row.initial, initial_merge_map),
                    final=resolve(row.final, final_merge_map),
                    tone_value=row.tone_value,
                )
            )
        unique_initials = self._sort_initials({row.initial for row in remapped_rows})
        unique_finals = self._sort_symbols({row.final for row in remapped_rows})
        unique_tones = self._sort_tones({row.tone_value for row in remapped_rows})
        unique_ipa = sorted({row.ipa for row in remapped_rows})
        return PhonologyAnalysisResult(
            rows=remapped_rows,
            unique_initials=unique_initials,
            unique_finals=unique_finals,
            unique_tones=unique_tones,
            unique_ipa=unique_ipa,
        )

    def export_outputs(
        self,
        analysis: PhonologyAnalysisResult,
        tone_class_map: dict[str, str],
        tone_value_order: list[str] | None,
        initial_order: list[str] | None,
        final_order: list[str] | None,
        output_dir: str,
    ) -> PhonologyOutputResult:
        out_dir = Path(output_dir)
        out_dir.mkdir(parents=True, exist_ok=True)
        forward_docx = out_dir / "同音字表_韵母到声母.docx"
        reverse_docx = out_dir / "同音字表_声母到韵母.docx"
        matrix_xlsx = out_dir / "同音字表_二维表.xlsx"
        ordered_initials = self._resolve_order(
            current=analysis.unique_initials,
            preferred=initial_order,
        )
        ordered_finals = self._resolve_order(
            current=analysis.unique_finals,
            preferred=final_order,
        )
        ordered_tones = self._resolve_order(
            current=analysis.unique_tones,
            preferred=tone_value_order,
        )
        self._write_word_document(
            analysis,
            tone_class_map,
            forward_docx,
            mode="final_initial",
            ordered_initials=ordered_initials,
            ordered_finals=ordered_finals,
            ordered_tones=ordered_tones,
        )
        self._write_word_document(
            analysis,
            tone_class_map,
            reverse_docx,
            mode="initial_final",
            ordered_initials=ordered_initials,
            ordered_finals=ordered_finals,
            ordered_tones=ordered_tones,
        )
        self._write_matrix_xlsx(
            analysis,
            tone_class_map,
            matrix_xlsx,
            ordered_initials=ordered_initials,
            ordered_finals=ordered_finals,
            ordered_tones=ordered_tones,
        )
        return PhonologyOutputResult(
            forward_docx_path=str(forward_docx),
            reverse_docx_path=str(reverse_docx),
            matrix_xlsx_path=str(matrix_xlsx),
        )

    def _load_rows_from_excel(self, path: Path, skip_first_row: bool) -> list[PhonologyInputRow]:
        df = pd.read_excel(path, header=None, dtype=str)
        return self._rows_from_frame(df, skip_first_row=skip_first_row)

    def _load_rows_from_csv(self, path: Path, skip_first_row: bool) -> list[PhonologyInputRow]:
        df = pd.read_csv(path, header=None, dtype=str)
        return self._rows_from_frame(df, skip_first_row=skip_first_row)

    def _load_rows_from_text(self, path: Path, skip_first_row: bool) -> list[PhonologyInputRow]:
        rows: list[PhonologyInputRow] = []
        with path.open("r", encoding="utf-8-sig", newline="") as f:
            for line_idx, line in enumerate(f):
                if skip_first_row and line_idx == 0:
                    continue
                clean_line = line.strip()
                if not clean_line:
                    continue
                cols = re.split(r"[\t,，]", clean_line, maxsplit=2)
                parsed = self._parse_columns(cols)
                if parsed is not None:
                    rows.append(parsed)
        return rows

    def _rows_from_frame(self, df: pd.DataFrame, skip_first_row: bool) -> list[PhonologyInputRow]:
        rows: list[PhonologyInputRow] = []
        for row_idx, values in df.iterrows():
            if skip_first_row and row_idx == 0:
                continue
            cols = ["" if pd.isna(v) else str(v) for v in values.tolist()]
            parsed = self._parse_columns(cols)
            if parsed is not None:
                rows.append(parsed)
        return rows

    def _parse_columns(self, cols: list[str]) -> PhonologyInputRow | None:
        if not cols:
            return None
        raw_char_col = cols[0].strip() if len(cols) >= 1 else ""
        ipa = cols[1].strip() if len(cols) >= 2 else ""
        note_col = cols[2].strip() if len(cols) >= 3 else ""
        if not raw_char_col and not ipa:
            return None
        if self._is_header_row(raw_char_col, ipa):
            return None
        char, note_from_char = self._extract_character_and_note(raw_char_col)
        note = self._merge_notes(note_from_char, note_col)
        if not char or not ipa:
            return None
        return PhonologyInputRow(character=char, ipa=ipa, note=note)

    def _extract_character_and_note(self, value: str) -> tuple[str, str]:
        raw = value.strip()
        if not raw:
            return "", ""
        chinese_chars = re.findall(r"[\u3400-\u9FFF\U00020000-\U0002A6DF]", raw)
        if not chinese_chars:
            return raw[:1], raw[1:].strip()
        char = chinese_chars[0]
        bracket_matches = re.findall(r"[（(]\s*([^）)]+?)\s*[）)]", raw)
        if bracket_matches:
            return char, "，".join(item.strip() for item in bracket_matches if item.strip())
        if len(chinese_chars) > 1:
            return char, "".join(chinese_chars)
        return char, ""

    def _merge_notes(self, note_a: str, note_b: str) -> str:
        values = [v.strip() for v in [note_a, note_b] if v and v.strip()]
        if not values:
            return ""
        seen: set[str] = set()
        merged: list[str] = []
        for item in values:
            if item in seen:
                continue
            seen.add(item)
            merged.append(item)
        return "，".join(merged)

    def _is_header_row(self, char_col: str, ipa_col: str) -> bool:
        c = char_col.replace(" ", "")
        p = ipa_col.replace(" ", "")
        header_chars = {"汉字", "字头", "字", "字符"}
        header_ipa = {"音标", "ipa", "拼音"}
        return c.lower() in {h.lower() for h in header_chars} and p.lower() in {
            h.lower() for h in header_ipa
        }

    def _sort_symbols(self, values: set[str]) -> list[str]:
        symbol_list = list(values)
        return sorted(symbol_list, key=lambda x: (x != "Ø", x == "", len(x), x))

    def _sort_initials(self, values: set[str]) -> list[str]:
        symbol_list = list(values)
        return sorted(symbol_list, key=self._initial_sort_key)

    def _sort_tones(self, tones: set[str]) -> list[str]:
        def key(value: str) -> tuple[int, int, str]:
            if value.isdigit():
                return (0, int(value), value)
            return (1, 0, value)

        return sorted(tones, key=key)

    def _resolve_order(self, current: list[str], preferred: list[str] | None) -> list[str]:
        if not preferred:
            return list(current)
        preferred_unique = []
        seen: set[str] = set()
        for item in preferred:
            if item in current and item not in seen:
                preferred_unique.append(item)
                seen.add(item)
        for item in current:
            if item not in seen:
                preferred_unique.append(item)
        return preferred_unique

    def _group_rows(
        self,
        analysis: PhonologyAnalysisResult,
        tone_class_map: dict[str, str],
        outer_key: str,
        inner_key: str,
    ) -> dict[str, dict[str, dict[str, list[ParsedPhonologyRow]]]]:
        grouped: dict[str, dict[str, dict[str, list[ParsedPhonologyRow]]]] = defaultdict(
            lambda: defaultdict(lambda: defaultdict(list))
        )
        for row in analysis.rows:
            outer = row.final if outer_key == "final" else row.initial
            inner = row.initial if inner_key == "initial" else row.final
            tone_label = tone_class_map.get(row.tone_value, row.tone_value or "0")
            grouped[outer][inner][tone_label].append(row)
        return grouped

    def _write_word_document(
        self,
        analysis: PhonologyAnalysisResult,
        tone_class_map: dict[str, str],
        output_path: Path,
        mode: str,
        ordered_initials: list[str],
        ordered_finals: list[str],
        ordered_tones: list[str],
    ) -> None:
        Document = self._resolve_document_factory()

        doc = Document()
        self._setup_word_document_styles(doc)
        title = doc.add_heading("同音字表", level=1)
        self._center_word_paragraph(title)
        if mode == "final_initial":
            subtitle = doc.add_paragraph("（韵母 → 声母）")
            grouped = self._group_rows(
                analysis=analysis,
                tone_class_map=tone_class_map,
                outer_key="final",
                inner_key="initial",
            )
            outer_values = ordered_finals
            inner_values = ordered_initials
        else:
            subtitle = doc.add_paragraph("（声母 → 韵母）")
            grouped = self._group_rows(
                analysis=analysis,
                tone_class_map=tone_class_map,
                outer_key="initial",
                inner_key="final",
            )
            outer_values = ordered_initials
            inner_values = ordered_finals
        self._center_word_paragraph(subtitle)
        self._style_word_paragraph_runs(subtitle, size=11, bold=True)

        self._append_summary_sections(
            doc,
            analysis,
            tone_class_map,
            ordered_initials=ordered_initials,
            ordered_finals=ordered_finals,
            ordered_tones=ordered_tones,
        )
        for outer in outer_values:
            if outer not in grouped:
                continue
            category_label = outer if outer else "空韵"
            category = doc.add_heading(f"{category_label}", level=3)
            self._center_word_paragraph(category)
            self._style_word_paragraph_runs(category, size=12, bold=True)
            inner_map = grouped[outer]
            for inner in inner_values:
                if inner not in inner_map:
                    continue
                p = doc.add_paragraph()
                self._add_word_text(p, f"{inner} ", size=12, bold=True)
                tone_map = inner_map[inner]
                for tone_label in self._ordered_tone_labels(
                    set(tone_map.keys()), tone_class_map, ordered_tones
                ):
                    self._add_word_text(p, f"[{tone_label}]")
                    self._append_word_entries(p, tone_map[tone_label])
            spacer = doc.add_paragraph("")
            self._style_word_paragraph_runs(spacer)
        doc.save(output_path)

    def _resolve_document_factory(self):
        try:
            from docx import Document as DocxDocument

            return DocxDocument
        except Exception:
            return _FallbackDocument

    def _append_summary_sections(
        self,
        doc,
        analysis: PhonologyAnalysisResult,
        tone_class_map: dict[str, str],
        ordered_initials: list[str],
        ordered_finals: list[str],
        ordered_tones: list[str],
    ) -> None:
        heading_initial = doc.add_heading("声母统计", level=2)
        self._center_word_paragraph(heading_initial)
        self._style_word_paragraph_runs(heading_initial, size=12, bold=True)
        initial_lines = []
        for initial in ordered_initials:
            examples = self._pick_examples(analysis.rows, lambda row: row.initial == initial)
            initial_lines.append((f"{initial}: ", examples))
        self._append_two_column_entries(doc, initial_lines)

        heading_final = doc.add_heading("韵母统计", level=2)
        self._center_word_paragraph(heading_final)
        self._style_word_paragraph_runs(heading_final, size=12, bold=True)
        final_lines = []
        for final in ordered_finals:
            label = final if final else "空韵"
            examples = self._pick_examples(analysis.rows, lambda row: row.final == final)
            final_lines.append((f"{label}: ", examples))
        self._append_two_column_entries(doc, final_lines)

        heading_tone = doc.add_heading("声调统计", level=2)
        self._center_word_paragraph(heading_tone)
        self._style_word_paragraph_runs(heading_tone, size=12, bold=True)
        tone_lines = []
        for tone_value in ordered_tones:
            tone_label = tone_class_map.get(tone_value, tone_value or "0")
            examples = self._pick_examples(
                analysis.rows, lambda row: row.tone_value == tone_value
            )
            tone_lines.append((f"{tone_value} → {tone_label}: ", examples))
        self._append_two_column_entries(doc, tone_lines)
        spacer = doc.add_paragraph("")
        self._style_word_paragraph_runs(spacer)

    def _pick_examples(
        self,
        rows: list[ParsedPhonologyRow],
        predicate,
        max_count: int = 5,
    ) -> list[ParsedPhonologyRow]:
        picked: list[ParsedPhonologyRow] = []
        seen: set[tuple[str, str]] = set()
        for row in rows:
            if not predicate(row):
                continue
            key = (row.character, row.note)
            if key in seen:
                continue
            seen.add(key)
            picked.append(row)
            if len(picked) >= max_count:
                break
        return picked

    def _append_word_entries(self, paragraph, entries: list[ParsedPhonologyRow]) -> None:
        for idx, row in enumerate(entries):
            if idx > 0:
                self._add_word_text(paragraph, " ")
            self._add_word_text(paragraph, row.character)
            if row.note:
                self._add_word_text(paragraph, row.note, size=12, subscript=True)

    def _write_matrix_xlsx(
        self,
        analysis: PhonologyAnalysisResult,
        tone_class_map: dict[str, str],
        output_path: Path,
        ordered_initials: list[str],
        ordered_finals: list[str],
        ordered_tones: list[str],
    ) -> None:
        grouped = self._group_rows(
            analysis=analysis,
            tone_class_map=tone_class_map,
            outer_key="final",
            inner_key="initial",
        )
        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "二维同音字表"
        header_fill = PatternFill("solid", fgColor="DCE6F1")
        thin_side = Side(style="thin", color="BFBFBF")
        thin_border = Border(left=thin_side, right=thin_side, top=thin_side, bottom=thin_side)
        center_alignment = Alignment(horizontal="center", vertical="center")
        body_alignment = Alignment(horizontal="left", vertical="top", wrap_text=True)

        sheet.cell(row=1, column=1, value="韵母\\声母")
        sheet.cell(row=1, column=1).font = Font(name="Times New Roman", size=11, bold=True)
        sheet.cell(row=1, column=1).alignment = center_alignment
        sheet.cell(row=1, column=1).fill = header_fill
        sheet.cell(row=1, column=1).border = thin_border
        for col_idx, initial in enumerate(ordered_initials, start=2):
            sheet.cell(row=1, column=col_idx, value=initial)
            sheet.cell(row=1, column=col_idx).font = Font(name="Times New Roman", size=11, bold=True)
            sheet.cell(row=1, column=col_idx).alignment = center_alignment
            sheet.cell(row=1, column=col_idx).fill = header_fill
            sheet.cell(row=1, column=col_idx).border = thin_border

        for row_idx, final in enumerate(ordered_finals, start=2):
            sheet.cell(row=row_idx, column=1, value=final if final else "")
            sheet.cell(row=row_idx, column=1).font = Font(name="Times New Roman", size=11, bold=True)
            sheet.cell(row=row_idx, column=1).alignment = center_alignment
            sheet.cell(row=row_idx, column=1).fill = header_fill
            sheet.cell(row=row_idx, column=1).border = thin_border
            inner_map = grouped.get(final, {})
            for col_idx, initial in enumerate(ordered_initials, start=2):
                tone_map = inner_map.get(initial, {})
                cell = sheet.cell(row=row_idx, column=col_idx)
                cell.alignment = body_alignment
                cell.border = thin_border
                cell.font = Font(name="Times New Roman", size=11)
                if not tone_map:
                    continue
                rich_value = self._build_rich_text_cell_value(tone_map)
                if rich_value is not None:
                    cell.value = rich_value
                else:
                    chunks: list[str] = []
                    for tone_label in self._ordered_tone_labels(
                        set(tone_map.keys()), tone_class_map, ordered_tones
                    ):
                        chars = "".join(
                            self._render_plain_entry(entry) for entry in tone_map[tone_label]
                        )
                        chunks.append(f"[{tone_label}]{chars}")
                    cell.value = " ".join(chunks)
        for col_idx in range(1, len(ordered_initials) + 2):
            col_letter = get_column_letter(col_idx)
            sheet.column_dimensions[col_letter].width = 18 if col_idx == 1 else 26
        for row_idx in range(1, len(ordered_finals) + 2):
            sheet.row_dimensions[row_idx].height = 36 if row_idx == 1 else 52
        sheet.freeze_panes = "B2"
        workbook.save(output_path)

    def _ordered_tone_labels(
        self,
        labels: set[str],
        tone_class_map: dict[str, str],
        ordered_tones: list[str],
    ) -> list[str]:
        ordered_labels: list[str] = []
        seen: set[str] = set()
        for tone_value in ordered_tones:
            label = tone_class_map.get(tone_value, tone_value or "0")
            if label in labels and label not in seen:
                ordered_labels.append(label)
                seen.add(label)
        for label in self._sort_tones(labels):
            if label not in seen:
                ordered_labels.append(label)
                seen.add(label)
        return ordered_labels

    def _append_two_column_entries(self, doc, entries: list[tuple[str, list[ParsedPhonologyRow]]]):
        if not hasattr(doc, "add_table"):
            for label, examples in entries:
                p = doc.add_paragraph()
                self._add_word_text(p, label, size=12, bold=True)
                self._append_word_entries(p, examples)
            return
        rows = (len(entries) + 1) // 2
        if rows == 0:
            return
        table = doc.add_table(rows=rows, cols=2)
        table.style = "Table Grid"
        for idx, (label, examples) in enumerate(entries):
            r = idx % rows
            c = idx // rows
            cell = table.cell(r, c)
            para = cell.paragraphs[0]
            self._add_word_text(para, label, size=11, bold=True)
            self._append_word_entries(para, examples)
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if not paragraph.text.strip():
                        paragraph.text = ""

    def _initial_sort_key(self, symbol: str):
        if symbol == "Ø":
            return (-1, -1, 0, symbol)
        if symbol == "":
            return (99, 99, 99, symbol)
        manner, place = self._classify_initial(symbol)
        manner_idx = self.INITIAL_MANNER_ORDER.index(manner) if manner in self.INITIAL_MANNER_ORDER else 99
        place_idx = self.INITIAL_PLACE_ORDER.index(place) if place in self.INITIAL_PLACE_ORDER else 99
        return (place_idx, manner_idx, len(symbol), symbol)

    def _classify_initial(self, symbol: str) -> tuple[str, str]:
        s = symbol
        if any(x in s for x in ["m", "n", "ŋ", "ɲ", "ɳ", "ɴ", "ȵ"]):
            manner = "鼻音"
        elif any(x in s for x in ["ts", "tɕ", "dʑ", "tʂ", "ɖʐ", "tʃ", "dʒ"]):
            manner = "塞擦音"
        elif any(x in s for x in ["p", "b", "t", "d", "k", "ɡ", "q", "ɢ", "ʔ", "ȶ", "ȡ"]):
            manner = "塞音"
        elif any(x in s for x in ["s", "z", "ʃ", "ʒ", "ʂ", "ʐ", "ɕ", "ʑ", "f", "v", "x", "ɣ", "h", "ɦ", "χ", "ʁ"]):
            manner = "擦音"
        elif any(x in s for x in ["l", "ɭ", "ʎ", "ʟ"]):
            manner = "边音"
        elif any(x in s for x in ["ɾ", "ɽ", "ɺ"]):
            manner = "闪音"
        elif any(x in s for x in ["r", "ʀ", "ʙ"]):
            manner = "颤音"
        elif any(x in s for x in ["ɹ", "ɻ", "j", "w", "ɰ", "ʋ"]):
            manner = "近音"
        else:
            manner = "其他"

        if any(x in s for x in ["p", "b", "m", "ʘ"]):
            place = "双唇"
        elif any(x in s for x in ["f", "v", "ɱ"]):
            place = "唇齿"
        elif any(x in s for x in ["θ", "ð"]):
            place = "舌尖前"
        elif any(x in s for x in ["t", "d", "n", "s", "z", "l", "ɾ", "r"]):
            place = "舌尖中"
        elif any(x in s for x in ["ʈ", "ɖ", "ɳ", "ʂ", "ʐ", "ɻ"]):
            place = "卷舌"
        elif any(x in s for x in ["ʃ", "ʒ"]):
            place = "龈后"
        elif any(x in s for x in ["ɕ", "ʑ", "ȶ", "ȡ", "ȵ", "tɕ", "dʑ"]):
            place = "龈腭"
        elif any(x in s for x in ["c", "ɟ", "ɲ", "j"]):
            place = "硬腭"
        elif any(x in s for x in ["k", "ɡ", "x", "ɣ", "ŋ", "w", "ɰ"]):
            place = "软腭"
        elif any(x in s for x in ["q", "ɢ", "χ", "ʁ", "ɴ"]):
            place = "小舌"
        elif any(x in s for x in ["ħ", "ʕ", "ʜ", "ʢ"]):
            place = "咽"
        elif any(x in s for x in ["h", "ɦ", "ʔ"]):
            place = "声门"
        else:
            place = "其他"
        return manner, place

    def _build_rich_text_cell_value(
        self, tone_map: dict[str, list[ParsedPhonologyRow]]
    ):
        if not _OPENPYXL_RICH_TEXT_AVAILABLE:
            return None
        rich_text = CellRichText()
        is_first_tone = True
        for tone_label in self._sort_tones(set(tone_map.keys())):
            if not is_first_tone:
                self._append_xlsx_rich_text(
                    rich_text, " ", font_name="Times New Roman", size=11
                )
            self._append_xlsx_rich_text(
                rich_text, f"[{tone_label}]", font_name="Times New Roman", size=11
            )
            for row in tone_map[tone_label]:
                self._append_xlsx_mixed_text(rich_text, row.character, size=11)
                if row.note:
                    self._append_xlsx_mixed_text(rich_text, row.note, size=11, subscript=True)
            is_first_tone = False
        return rich_text

    def _append_xlsx_mixed_text(self, rich_text, text: str, size: int, subscript: bool = False):
        if not text:
            return
        current_buffer = ""
        current_cjk = self._contains_cjk(text[0])
        for ch in text:
            flag = self._contains_cjk(ch)
            if flag != current_cjk and current_buffer:
                self._append_xlsx_rich_text(
                    rich_text,
                    current_buffer,
                    font_name="宋体" if current_cjk else "Times New Roman",
                    size=size,
                    subscript=subscript,
                )
                current_buffer = ""
                current_cjk = flag
            current_buffer += ch
        if current_buffer:
            self._append_xlsx_rich_text(
                rich_text,
                current_buffer,
                font_name="宋体" if current_cjk else "Times New Roman",
                size=size,
                subscript=subscript,
            )

    def _append_xlsx_rich_text(
        self,
        rich_text,
        text: str,
        font_name: str,
        size: int,
        subscript: bool = False,
    ):
        if not text:
            return
        font = InlineFont(
            rFont=font_name,
            sz=size,
            vertAlign="subscript" if subscript else None,
        )
        rich_text.append(TextBlock(font, text))

    def _setup_word_document_styles(self, doc):
        if not hasattr(doc, "styles"):
            return
        try:
            from docx.oxml.ns import qn
            from docx.shared import Pt
        except Exception:
            return
        style = doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(12)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _center_word_paragraph(self, paragraph):
        if not hasattr(paragraph, "alignment"):
            return
        try:
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        except Exception:
            return
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    def _style_word_paragraph_runs(self, paragraph, size: int = 12, bold: bool = False):
        runs = getattr(paragraph, "runs", [])
        for run in runs:
            self._apply_word_run_style(run, size=size, bold=bold)

    def _add_word_text(
        self,
        paragraph,
        text: str,
        size: int = 12,
        bold: bool = False,
        subscript: bool = False,
    ):
        if not text:
            return
        if not hasattr(paragraph, "runs"):
            run = paragraph.add_run(text)
            run.bold = bold
            if subscript:
                run.font.subscript = True
            return
        current_buffer = ""
        current_cjk = self._contains_cjk(text[0])
        for ch in text:
            flag = self._contains_cjk(ch)
            if flag != current_cjk and current_buffer:
                run = paragraph.add_run(current_buffer)
                self._apply_word_run_style(
                    run,
                    size=size,
                    bold=bold,
                    subscript=subscript,
                    cjk=current_cjk,
                )
                current_buffer = ""
                current_cjk = flag
            current_buffer += ch
        if current_buffer:
            run = paragraph.add_run(current_buffer)
            self._apply_word_run_style(
                run,
                size=size,
                bold=bold,
                subscript=subscript,
                cjk=current_cjk,
            )

    def _apply_word_run_style(
        self,
        run,
        size: int = 12,
        bold: bool = False,
        subscript: bool = False,
        cjk: bool = False,
    ):
        if not hasattr(run, "font"):
            return
        try:
            from docx.oxml.ns import qn
            from docx.shared import Pt
        except Exception:
            return
        run.bold = bold
        run.font.size = Pt(size)
        run.font.subscript = subscript
        run.font.name = "宋体" if cjk else "Times New Roman"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")

    def _contains_cjk(self, text: str) -> bool:
        for ch in text:
            code = ord(ch)
            if 0x4E00 <= code <= 0x9FFF:
                return True
            if 0x3400 <= code <= 0x4DBF:
                return True
            if 0x20000 <= code <= 0x2A6DF:
                return True
        return False

    def _render_plain_entry(self, row: ParsedPhonologyRow) -> str:
        if not row.note:
            return row.character
        return f"{row.character}{self._to_subscript_text(row.note)}"

    def _to_subscript_text(self, value: str) -> str:
        mapping = str.maketrans(
            {
                "0": "₀",
                "1": "₁",
                "2": "₂",
                "3": "₃",
                "4": "₄",
                "5": "₅",
                "6": "₆",
                "7": "₇",
                "8": "₈",
                "9": "₉",
                "(": "₍",
                ")": "₎",
                "+": "₊",
                "-": "₋",
                "=": "₌",
                "a": "ₐ",
                "e": "ₑ",
                "o": "ₒ",
                "x": "ₓ",
                "ə": "ₔ",
            }
        )
        return value.translate(mapping)
